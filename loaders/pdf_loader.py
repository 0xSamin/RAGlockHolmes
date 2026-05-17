from pypdf import PdfReader
from docx import Document
from typing import List, Dict, Any
import os

class BaseLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def read(self):
        raise NotImplementedError("Subclasses must implement the 'read' method.")


class PdfLoader(BaseLoader):
    def extract_metadata(self):
        """Extract PDF metadata (/Author, /Title, /Subject, /CreationDate)"""
        try:
            reader = PdfReader(self.file_path)
            metadata = reader.metadata if reader.metadata else {}

            return {
                'author': metadata.get('/Author', 'Unknown'),
                'title': metadata.get('/Title', 'Unknown'),
                'subject': metadata.get('/Subject', ''),
                'creation_date': metadata.get('/CreationDate', '')
            }
        except Exception as e:
            print(f"Warning: Could not extract metadata from {self.file_path}: {e}")
            return {
                'author': 'Unknown',
                'title': 'Unknown',
                'subject': '',
                'creation_date': ''
            }

    def read(self):
        """
        Extract text from PDF with metadata.
        Returns list starting with synthetic page 0 (metadata + first page excerpt),
        followed by regular pages.
        """
        paper = PdfReader(self.file_path)

        # Extract metadata
        doc_metadata = self.extract_metadata()

        # Extract first page text for synthetic chunk
        first_page_text = ""
        if len(paper.pages) > 0:
            first_page_text = paper.pages[0].extract_text()
            if first_page_text:
                first_page_text = first_page_text.strip()[:500]  # First 500 chars

        # Create synthetic page 0 chunk (metadata + excerpt)
        synthetic_text = (
            f"Document Title: {doc_metadata['title']}\n"
            f"Author: {doc_metadata['author']}\n"
            f"Subject: {doc_metadata['subject']}\n\n"
            f"First Page Excerpt:\n{first_page_text}"
        )

        text_data = [{
            "page_number": 0,
            "text": synthetic_text,
            "metadata": doc_metadata,
            "is_synthetic": True
        }]

        # Extract regular pages
        for i, page in enumerate(paper.pages):
            extracted_text = page.extract_text()
            if extracted_text and extracted_text.strip():
                text_data.append({
                    "page_number": i + 1,
                    "text": extracted_text.strip(),
                    "metadata": doc_metadata,
                    "is_synthetic": False
                })

        return text_data


class DocxLoader:
    """Loader for Word documents (.docx) with metadata extraction."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

    def extract_metadata(self) -> Dict[str, str]:
        """Extract DOCX metadata from core properties."""
        try:
            doc = Document(self.file_path)
            props = doc.core_properties

            return {
                'author': props.author or 'Unknown',
                'title': props.title or os.path.basename(self.file_path),
                'subject': props.subject or '',
                'creation_date': str(props.created) if props.created else ''
            }
        except Exception as e:
            print(f"Warning: Could not extract metadata from {self.file_path}: {e}")
            return {
                'author': 'Unknown',
                'title': os.path.basename(self.file_path),
                'subject': '',
                'creation_date': ''
            }

    def read(self) -> List[Dict[str, Any]]:
        """
        Extract text from Word document with metadata.
        Returns list starting with synthetic page 0 (metadata + first page excerpt),
        followed by regular pages.
        """
        doc = Document(self.file_path)

        # Extract metadata
        doc_metadata = self.extract_metadata()

        # Collect all non-empty paragraphs
        all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        if not all_paragraphs:
            return []

        # Extract first ~500 characters for synthetic chunk
        first_page_text = "\n".join(all_paragraphs[:5])[:500]

        # Create synthetic page 0 chunk (metadata + excerpt)
        synthetic_text = (
            f"Document Title: {doc_metadata['title']}\n"
            f"Author: {doc_metadata['author']}\n"
            f"Subject: {doc_metadata['subject']}\n\n"
            f"First Page Excerpt:\n{first_page_text}"
        )

        text_data = [{
            "page_number": 0,
            "text": synthetic_text,
            "metadata": doc_metadata,
            "is_synthetic": True
        }]

        # Split paragraphs into "pages" (every ~10 paragraphs = 1 page)
        paragraphs_per_page = 5
        page_number = 1

        for i in range(0, len(all_paragraphs), paragraphs_per_page):
            page_text = "\n".join(all_paragraphs[i:i + paragraphs_per_page])

            if page_text.strip():
                text_data.append({
                    "page_number": page_number,
                    "text": page_text.strip(),
                    "metadata": doc_metadata,
                    "is_synthetic": False
                })
                page_number += 1

        return text_data
